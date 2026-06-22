#!/usr/bin/env python3
"""Generate Noah Michaels' cover letter for the WA Correctional Industries
Brand & Communications Designer (Communications Consultant 4) role.

Writes NoahMichaelsCoverLetter-CI.docx set in Archivo, then exports a PDF via
LibreOffice (soffice). Run with the local venv:

    cover-letters/.venv/bin/python cover-letters/make-cover-letter.py
"""

from __future__ import annotations

import datetime
import re
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

HERE = Path(__file__).resolve().parent
DOCX_PATH = HERE / "NoahMichaelsCoverLetter-CI.docx"

FONT = "Archivo"
BODY_PT = 11
NAME_PT = 20
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x55, 0x55, 0x55)

CONTACT = "919.376.6669  |  hi@nm.works  |  nm.works"
LOCATION = "Raleigh, NC \u00b7 Open to relocation"

GREETING = "Dear Correctional Industries Hiring Team,"

BODY = [
    "When the North Carolina General Assembly needed an identity for its America "
    "250 (Semiquincentennial) Committee, I designed the official seal and the brand "
    "system that carried it across print and digital. Owning a brand from concept "
    "to completion and keeping it consistent everywhere it appears is exactly what "
    "this role calls for, and it is what I'd bring to Correctional Industries as "
    "your Brand & Communications Designer (Communications Consultant 4).",
    "For five years I've designed in the Adobe Creative Suite, working daily in "
    "InDesign, Illustrator, and Photoshop to produce the full range of materials "
    "your team handles: logos, catalogs, reports, flyers, signage, and event "
    "collateral. I build the template systems that hold brand consistency across "
    "that output, and I've carried it across 20+ responsive websites in WordPress "
    "and Drupal. Because written content carries the message as much as the visuals "
    "do, I draft, edit, and proofread every piece as both.",
    "What sets me apart is that I know how the work is produced. Facing personalized "
    "mailings to all 170 members of the General Assembly that normally took two "
    "passes through the press, I wrote a variable-data merge combining base "
    "stationery with per-recipient content in a single pass, cutting the run in "
    "half. I run tens of thousands of impressions weekly during legislative "
    "session, handling imposition, pre-flight, and color on EFI Fiery, and I've "
    "worked the commercial-print side from detailed publication specifications and "
    "cost estimates through press checks, the preferred experience this posting "
    "names.",
    "That print shop is a fast-paced, deadline-driven environment where I prioritize "
    "competing projects, meet deadlines, and work independently while collaborating "
    "with a small team. I'd partner closely with your Webmaster, since I also build "
    "front-ends in HTML, CSS, and JavaScript, and I hold a B.S. in Graphic Arts & "
    "Imaging Technology. CI's commitment to integrity, inclusion, and supporting "
    "people's success resonates with me, and I'm ready to relocate to Tumwater.",
    "I'd welcome the chance to talk through how I can help advance the CI brand and "
    "the mission behind it. Thank you for your consideration.",
]

AGENCY = [
    "Washington State Department of Corrections",
    "Correctional Industries, Marketing & Communications",
    "Tumwater, WA",
]


def smarten(text: str) -> str:
    """Convert straight quotes/apostrophes to typographic (curly) ones."""
    text = re.sub(r"(^|[\s(\[{\u2014])\"", "\\1\u201c", text)  # opening double
    text = text.replace('"', "\u201d")  # closing double
    text = re.sub(r"(^|[\s(\[{\u2014])'", "\\1\u2018", text)  # opening single
    text = text.replace("'", "\u2019")  # apostrophe / closing single
    return text


def _set_run(run, *, size=BODY_PT, bold=False, color=INK):
    run.text = smarten(run.text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), FONT)


def _para(doc, *, before=0.0, after=10.0, line=1.18, align=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        pf.alignment = align
    return p


def build():
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(BODY_PT)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    for section in doc.sections:
        section.top_margin = Pt(54)
        section.bottom_margin = Pt(54)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

    # Header: name + contact
    p = _para(doc, after=2.0)
    _set_run(p.add_run("Noah Michaels"), size=NAME_PT, bold=True)

    p = _para(doc, after=1.0)
    _set_run(p.add_run(CONTACT), size=9.5, color=MUTED)

    p = _para(doc, after=14.0)
    _set_run(p.add_run(LOCATION), size=9.5, color=MUTED)

    # Date
    today = datetime.date.today().strftime("%B %-d, %Y")
    p = _para(doc, after=11.0)
    _set_run(p.add_run(today))

    # Agency address block
    for i, line in enumerate(AGENCY):
        p = _para(doc, after=(11.0 if i == len(AGENCY) - 1 else 1.0))
        _set_run(p.add_run(line))

    # Greeting
    p = _para(doc, after=10.0)
    _set_run(p.add_run(GREETING))

    # Body
    for para in BODY:
        p = _para(doc, after=8.0)
        _set_run(p.add_run(para))

    # Signature
    p = _para(doc, before=6.0, after=1.0)
    _set_run(p.add_run("Sincerely,"))
    p = _para(doc, after=0.0)
    _set_run(p.add_run("Noah Michaels"), bold=True)

    doc.save(DOCX_PATH)
    print(f"Wrote {DOCX_PATH}")


def export_pdf():
    soffice = shutil.which("soffice")
    if not soffice:
        print("soffice not found; skipping PDF export")
        return
    subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(HERE), str(DOCX_PATH)],
        check=True,
    )
    print(f"Wrote {DOCX_PATH.with_suffix('.pdf')}")


if __name__ == "__main__":
    build()
    export_pdf()
