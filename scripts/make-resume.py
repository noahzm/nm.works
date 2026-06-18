from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

BODY = "Univers Next Pro"
INK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x55, 0x55, 0x55)
RIGHT_EDGE = Inches(7.0)
REPO_ROOT = Path(__file__).resolve().parents[1]

doc = Document()

sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.75)
sec.bottom_margin = Inches(0.75)
sec.left_margin = Inches(0.75)
sec.right_margin = Inches(0.75)

normal = doc.styles["Normal"]
normal.font.name = BODY
normal.font.size = Pt(10)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
normal.paragraph_format.line_spacing = 1.15


def tracking(run, twentieths):
    rPr = run._element.get_or_add_rPr()
    el = OxmlElement("w:spacing")
    el.set(qn("w:val"), str(twentieths))
    rPr.append(el)


def bottom_rule(par, size=4, color="1A1A1A", space=4):
    pPr = par._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def para(before=0, after=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    return p


def run(p, text, size=10, bold=False, italic=False, color=INK, font=BODY):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = color
    return r


def section(title, before=14):
    p = para(before=before, after=6)
    r = run(p, title.upper(), size=9, bold=True)
    tracking(r, 30)
    bottom_rule(p)
    return p


def heading_line(left_bold, left_plain, right_text, before=8):
    p = para(before=before, after=2)
    p.paragraph_format.tab_stops.add_tab_stop(RIGHT_EDGE, WD_TAB_ALIGNMENT.RIGHT)
    run(p, left_bold, size=10.5, bold=True)
    if left_plain:
        run(p, left_plain, size=10.5)
    run(p, "\t" + right_text, size=10, color=GRAY)
    return p


def hyperlink(p, url, text, size=9.5, color="555555", italic=False):
    r_id = p.part.relate_to(url, RT.HYPERLINK, is_external=True)
    h = OxmlElement("w:hyperlink")
    h.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), BODY)
    rFonts.set(qn("w:hAnsi"), BODY)
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)
    col = OxmlElement("w:color")
    col.set(qn("w:val"), color)
    rPr.append(col)
    if italic:
        rPr.append(OxmlElement("w:i"))
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    h.append(r)
    p._p.append(h)


def bullet(text):
    p = para(after=2)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    run(p, "•  ", color=GRAY)
    run(p, text)
    return p


# ---- Header ----
p = para(after=1)
r = run(p, "Noah Michaels", size=21, bold=True)
tracking(r, 10)

p = para(after=4)
r = run(p, "UX ENGINEER", size=10, color=GRAY)
tracking(r, 40)

p = para(after=2)
run(p, "Raleigh, NC 27603   ·   919-376-6669   ·   ", size=9.5, color=GRAY)
hyperlink(p, "mailto:hi@nm.works", "hi@nm.works")
run(p, "   ·   ", size=9.5, color=GRAY)
hyperlink(p, "https://nm.works", "nm.works")
run(p, "   ·   ", size=9.5, color=GRAY)
hyperlink(p, "https://github.com/noahzm", "github.com/noahzm")
bottom_rule(p, size=8, space=6)

# ---- Profile ----
section("Profile", before=12)
p = para(after=0)
run(
    p,
    "UX engineer who designs in code across iOS and web. Six years making "
    "legislative template systems, print-shop ordering, and a cycling weather "
    "app that turns forecasts into a plain ride verdict. Uses AI-assisted "
    "coding tools daily for faster iteration; product and accessibility "
    "decisions stay author-led.",
)

# ---- Projects ----
section("Selected Product Work")

heading_line(
    "Wheely Weather",
    " · Native iOS cycling weather app",
    "2025",
    before=2,
)
p = para(after=2)
hyperlink(p, "https://wheelyweather.app", "wheelyweather.app", italic=True)
run(p, "  ·  Swift, SwiftUI · Astro, React, TypeScript", size=9.5, italic=True, color=GRAY)
bullet(
    "Designed and shipped a native iOS app that answers the question cyclists "
    "actually have (ride now, or wait?), with verdict-first UI, hourly "
    "windows, best-day picks, and kit guidance."
)
bullet(
    "Validated the product model through informal conversations with cyclists; "
    "shared forecast logic lives in WheelyCore with accessibility-aware UI."
)
bullet(
    "Earlier web version on Cloudflare Pages; watchOS companion via "
    "WatchConnectivity. App Store submission in progress."
)

p = para(before=4)
run(p, "Full case studies at ", size=9.5, italic=True, color=GRAY)
hyperlink(p, "https://nm.works/projects", "nm.works/projects", italic=True)

# ---- Experience ----
section("Experience")

heading_line(
    "Graphic & Production Designer",
    ", NC General Assembly · Raleigh, NC",
    "2021–Present",
    before=2,
)
bullet(
    "Built template systems for the cards, postcards, letterhead, and "
    "envelopes that 170+ legislators order on repeat, reducing rework so each "
    "job starts from a standard instead of a blank page."
)
bullet(
    "Replaced a two-step mailing process by data-merging letters directly "
    "onto letterhead, so a full mailing prints in one run instead of two."
)
bullet(
    "Reworked the structure and layout of data-heavy legislative "
    "publications so they’re easier to read, navigate, and produce "
    "accurately."
)

heading_line(
    "Web & Graphic Designer",
    ", Creative Printing · Boone, NC",
    "2018–2020",
)
bullet(
    "Redesigned online ordering around the jobs customers came to do, with "
    "a clearer path from picking a service to submitting the work."
)
bullet(
    "Designed and built 20+ WordPress sites, working directly with clients "
    "from first wireframes through scoping and rounds of revisions."
)

# ---- Education ----
section("Education")
heading_line("Appalachian State University", "", "2020", before=2)
p = para(after=0)
run(
    p,
    "B.S., Graphic Arts and Imaging Technology, Cross Media Production "
    "· Minor: General Business",
)

# ---- Skills ----
section("Skills")


def skill_line(label, items, before=2):
    p = para(before=before, after=2)
    run(p, label + ":  ", bold=True)
    run(p, items)


skill_line(
    "Product & UX Design",
    "user flows, information architecture, wireframing, prototyping, "
    "design systems, accessibility-aware design · Figma",
)
skill_line(
    "Front-End Development",
    "Swift, SwiftUI, React, TypeScript, Astro, Tailwind CSS, HTML/CSS · "
    "Git, GitHub, Cloudflare Pages",
)
skill_line(
    "AI-Assisted Development",
    "Cursor, Claude Code, Codex, OpenCode — agentic design-to-production "
    "workflow",
)
skill_line(
    "Visual Design",
    "typography, layout, brand and print production · Adobe Creative Cloud",
)

props = doc.core_properties
props.title = "Noah Michaels – UX Engineer – Resume"
props.author = "Noah Michaels"
props.subject = "Resume"

doc.save(REPO_ROOT / "NoahMichaelsResume.docx")
print("saved", REPO_ROOT / "NoahMichaelsResume.docx")
